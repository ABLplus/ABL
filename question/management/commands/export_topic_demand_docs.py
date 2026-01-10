import os
import re

from django.core.management.base import BaseCommand, CommandError
from django.utils.text import slugify

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from syllabus.models import Subject, Section, Topic, TopicDemand


# ----------------------------
# Markdown helpers (basic)
# ----------------------------

def add_inline_md(paragraph, text):
    """
    Handle **bold**, *italic*, `code`
    """
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)

    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(tok)


def render_markdown(doc: Document, md: str):
    """
    Very practical Markdown renderer for DOCX:
    - #, ##, ###
    - bullets
    - numbered lists
    - code blocks
    - paragraphs
    """
    if not md:
        return

    lines = md.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # code block
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run(lines[i])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                i += 1
            i += 1
            continue

        # headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            h = doc.add_heading("", level=min(level, 3))
            add_inline_md(h, text)
            i += 1
            continue

        # bullet
        m_b = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m_b:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_md(p, m_b.group(1))
            i += 1
            continue

        # numbered
        m_n = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m_n:
            p = doc.add_paragraph(style="List Number")
            add_inline_md(p, m_n.group(1))
            i += 1
            continue

        # empty
        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        # paragraph
        p = doc.add_paragraph()
        add_inline_md(p, line)
        i += 1


# ----------------------------
# Command
# ----------------------------

class Command(BaseCommand):
    help = "Export TopicDemand Markdown as DOCX files (one per Section)."

    def add_arguments(self, parser):
        parser.add_argument("--subject-id", type=int, required=True)
        parser.add_argument(
            "--exam-name",
            type=str,
            default="UPSC CSE (Prelims)",
        )
        parser.add_argument(
            "--outdir",
            type=str,
            default="exports/topic_demand_docs",
        )
        parser.add_argument(
            "--include-missing",
            action="store_true",
        )

    def handle(self, *args, **opts):
        subject_id = opts["subject_id"]
        exam_name = opts["exam_name"]
        outdir = opts["outdir"]
        include_missing = opts["include_missing"]

        subject = Subject.objects.filter(id=subject_id).select_related("exam").first()
        if not subject:
            raise CommandError("Subject not found")

        if not os.path.isabs(outdir):
            outdir = os.path.join(os.getcwd(), outdir)
        os.makedirs(outdir, exist_ok=True)

        demands = TopicDemand.objects.filter(
            subject_id=subject_id,
            exam_name=exam_name,
        ).select_related("topic", "section")

        demand_by_topic = {d.topic_id: d for d in demands}

        sections = Section.objects.filter(subject=subject).prefetch_related("topics")

        for section in sections:
            doc = Document()

            # Title
            title = doc.add_heading(subject.name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_heading(section.name, level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            meta = doc.add_paragraph(f"Exam: {exam_name}")
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

            idx = 1
            for topic in section.topics.all():
                demand = demand_by_topic.get(topic.id)

                if not demand and not include_missing:
                    continue

                h = doc.add_heading(f"{idx}. {topic.name}", level=2)

                meta_p = doc.add_paragraph(
                    f"Tier: {topic.tier} | "
                    f"Weightage: {topic.weightage} | "
                    f"Total Qs: {topic.total_questions}"
                )
                meta_p.runs[0].italic = True

                if demand:
                    render_markdown(doc, demand.demand_insights)
                else:
                    doc.add_paragraph("No insights available.")

                idx += 1

            fname = f"{slugify(subject.name)}__{slugify(section.name)}__{slugify(exam_name)}.docx"
            path = os.path.join(outdir, fname)
            doc.save(path)

            self.stdout.write(self.style.SUCCESS(f"Generated: {path}"))

        self.stdout.write(self.style.SUCCESS("All DOCX exports completed."))